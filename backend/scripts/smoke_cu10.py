"""Run the real CU10 primary and fallback smoke tests against the Docker stack."""

import json
import sys
import time
import uuid
from io import BytesIO

import httpx
from docx import Document as DocxDocument
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.integrations.embeddings.dependencies import get_embedding_provider
from app.integrations.llm import (
    GeneratedQuestionBank,
    LLMErrorKind,
    QuestionGenerationProviderError,
    QuestionGenerationRequest,
)
from app.integrations.llm.dependencies import get_question_generation_router
from app.integrations.llm.router import QuestionGenerationRouter
from app.integrations.vector_db.dependencies import get_vector_store_provider
from app.modules.p1_gestion_identidad_seguridad.models.user import User
from app.modules.p2_gestion_documentos_preparacion.services.question_bank_service import (
    QuestionBankService,
)

API_URL = "http://localhost:8000/api/v1"


class ForcedGeminiFailure:
    name = "gemini"
    model = "forced-smoke-failure"

    def generate(self, _: QuestionGenerationRequest) -> GeneratedQuestionBank:
        raise QuestionGenerationProviderError(
            "Forced primary failure for the CU10 smoke test",
            provider=self.name,
            kind=LLMErrorKind.TRANSIENT,
            detail="forced_smoke",
        )


def build_document() -> bytes:
    sections = (
        (
            "Objetivos y problema",
            "Socratia busca preparar estudiantes para una defensa academica mediante preguntas "
            "sustentadas en sus propios documentos. El problema central es la falta de practica "
            "personalizada y de retroalimentacion trazable antes de presentarse ante un tribunal.",
        ),
        (
            "Metodologia",
            "La metodologia separa la gestion de identidad, la preparacion documental y la "
            "simulacion. Cada capacidad se valida con pruebas automatizadas y pruebas de humo "
            "reales. La recuperacion usa seis intenciones para cubrir distintas secciones.",
        ),
        (
            "Arquitectura tecnica",
            "El backend usa FastAPI, SQLAlchemy y PostgreSQL dentro de un monolito modular. "
            "Gemini genera embeddings, Pinecone recupera fragmentos y un router envia la "
            "generacion estructurada primero a Gemini y luego a Groq como respaldo.",
        ),
        (
            "Resultados y validacion",
            "La validacion exige doce preguntas: tres conceptuales, tres metodologicas, tres "
            "tecnicas y tres criticas. Las respuestas se validan con Pydantic y cada pregunta "
            "conserva referencias a los fragmentos que la sustentan.",
        ),
        (
            "Limitaciones y riesgos",
            "La calidad depende del texto extraido y de la cobertura de los fragmentos. Existen "
            "riesgos de latencia, cuotas y respuestas invalidas de proveedores externos. Los "
            "timeouts, reintentos acotados, fallback y circuit breaker reducen esos riesgos.",
        ),
        (
            "Conclusiones y trabajo futuro",
            "El banco persistido puede reutilizarse sin llamar nuevamente al modelo. En etapas "
            "futuras las preguntas alimentaran la simulacion oral y la evaluacion de respuestas "
            "contra puntos esperados que nunca se muestran al estudiante.",
        ),
    )
    document = DocxDocument()
    document.add_heading("Socratia: preparacion de defensas academicas", level=1)
    for heading, paragraph in sections:
        document.add_heading(heading, level=2)
        for _ in range(8):
            document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def assert_public_bank(bank: dict[str, object], expected_provider: str) -> None:
    questions = bank.get("questions")
    if not isinstance(questions, list) or len(questions) != 12:
        raise AssertionError("CU10 did not return exactly 12 questions")
    categories = {
        category: 0
        for category in ("CONCEPTUAL", "METHODOLOGICAL", "TECHNICAL", "CRITICAL")
    }
    for question in questions:
        if not isinstance(question, dict):
            raise AssertionError("CU10 returned an invalid public question")
        if set(question) != {"id", "question", "category", "difficulty"}:
            raise AssertionError("CU10 exposed fields outside the public question contract")
        categories[str(question["category"])] += 1
    if any(count != 3 for count in categories.values()):
        raise AssertionError("CU10 returned an invalid category distribution")
    if bank.get("provider_used") != expected_provider:
        raise AssertionError(f"CU10 expected provider {expected_provider}")


def api_primary_smoke() -> tuple[httpx.Client, str, str, dict[str, object], int]:
    client = httpx.Client(base_url=API_URL, timeout=300)
    suffix = uuid.uuid4().hex[:12]
    response = client.post(
        "/auth/register",
        json={
            "email": f"cu10-smoke-{suffix}@example.com",
            "full_name": "CU10 Real Smoke",
            "password": f"Socratia smoke phrase {uuid.uuid4().hex}",
        },
    )
    response.raise_for_status()
    user_id = response.json()["user"]["id"]

    response = client.post(
        "/documents",
        files={
            "file": (
                "socratia-cu10-smoke.docx",
                build_document(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    response.raise_for_status()
    document_id = response.json()["id"]

    response = client.post(f"/documents/{document_id}/process")
    response.raise_for_status()
    chunk_count = int(response.json()["chunk_count"])

    last_response: httpx.Response | None = None
    for attempt in range(3):
        if attempt:
            time.sleep(4)
        last_response = client.post(f"/documents/{document_id}/questions/generate")
        if last_response.is_success:
            break
    assert last_response is not None
    last_response.raise_for_status()
    bank = last_response.json()
    assert_public_bank(bank, "gemini")
    return client, user_id, document_id, bank, chunk_count


def real_fallback_smoke(user_id: str, document_id: str) -> None:
    settings = get_settings()
    engine = create_engine(settings.database_url, pool_pre_ping=True)
    configured_router = get_question_generation_router()
    router = QuestionGenerationRouter(
        primary=ForcedGeminiFailure(),
        fallback=configured_router.fallback,
        failure_threshold=1,
        recovery_seconds=60,
    )
    with Session(engine, expire_on_commit=False) as db:
        user = db.get(User, user_id)
        if user is None:
            raise AssertionError("Smoke user was not persisted")
        bank = QuestionBankService(
            db=db,
            embeddings=get_embedding_provider(),
            vectors=get_vector_store_provider(),
            generation_router=router,
        ).generate(user, document_id)
        if bank.provider_used != "groq" or not bank.fallback_used:
            raise AssertionError("Real Groq fallback was not persisted")


def main() -> int:
    client, user_id, document_id, primary_bank, chunk_count = api_primary_smoke()
    real_fallback_smoke(user_id, document_id)
    fallback_response = client.get(f"/documents/{document_id}/questions")
    fallback_response.raise_for_status()
    fallback_bank = fallback_response.json()
    assert_public_bank(fallback_bank, "groq")
    if fallback_bank.get("fallback_used") is not True:
        raise AssertionError("Fallback telemetry was not exposed")
    print(
        json.dumps(
            {
                "document_id": document_id,
                "chunk_count": chunk_count,
                "question_count": len(fallback_bank["questions"]),
                "primary_provider": primary_bank["provider_used"],
                "fallback_provider": fallback_bank["provider_used"],
                "fallback_used": fallback_bank["fallback_used"],
                "status": fallback_bank["status"],
            },
            ensure_ascii=False,
        )
    )
    client.close()
    return 0


if __name__ == "__main__":
    sys.exit(main())
