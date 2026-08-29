import re
import unicodedata
from io import BytesIO

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError


class TextExtractionError(Exception):
    """Raised when a supported document cannot yield usable text."""


class TextExtractionService:
    def extract(self, *, file_type: str, content: bytes) -> str:
        if file_type == "PDF":
            text = self._extract_pdf(content)
        elif file_type == "DOCX":
            text = self._extract_docx(content)
        else:
            raise TextExtractionError("Unsupported document type")

        normalized = self.normalize(text)
        if not normalized:
            raise TextExtractionError("The document does not contain extractable text")
        return normalized

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(content))
            if reader.is_encrypted:
                try:
                    unlocked = reader.decrypt("")
                except Exception as exc:  # pypdf exposes several encryption backends
                    raise TextExtractionError("The PDF is password protected") from exc
                if unlocked == 0:
                    raise TextExtractionError("The PDF is password protected")
            pages = [page.extract_text() or "" for page in reader.pages]
        except TextExtractionError:
            raise
        except (PdfReadError, ValueError, OSError) as exc:
            raise TextExtractionError("The PDF could not be read") from exc
        except Exception as exc:
            raise TextExtractionError("The PDF text could not be extracted") from exc
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        try:
            document = DocxDocument(BytesIO(content))
        except (PackageNotFoundError, ValueError, KeyError, OSError) as exc:
            raise TextExtractionError("The DOCX could not be read") from exc
        except Exception as exc:
            raise TextExtractionError("The DOCX text could not be extracted") from exc

        blocks: list[str] = []
        blocks.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)

    @staticmethod
    def normalize(text: str) -> str:
        text = unicodedata.normalize("NFC", text).replace("\x00", "")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
        lines = [re.sub(r"[^\S\n]+", " ", line).strip() for line in text.split("\n")]
        normalized_lines: list[str] = []
        previous_blank = True
        for line in lines:
            is_blank = not line
            if is_blank and previous_blank:
                continue
            normalized_lines.append(line)
            previous_blank = is_blank
        return "\n".join(normalized_lines).strip()
