from io import BytesIO
from typing import Any

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.integrations.llm import (
    LLMErrorKind,
    QuestionContextChunk,
    QuestionGenerationProviderError,
    QuestionGenerationRequest,
)
from app.integrations.llm.router import QuestionGenerationRouter
from app.modules.p2_gestion_documentos_preparacion.models.question_bank import (
    Question,
    QuestionBank,
    QuestionBankStatus,
)


def valid_docx() -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_heading("Sistema Socratia", level=1)
    document.add_paragraph(
        (
            "El objetivo es preparar defensas academicas mediante recuperacion semantica. "
            "La metodologia separa documentos por usuario y valida resultados con pruebas. "
            "La arquitectura usa servicios desacoplados, PostgreSQL y una base vectorial. "
            "Los resultados muestran preguntas sustentadas en evidencia. "
            "Las limitaciones incluyen la calidad del documento y la latencia de proveedores. "
            "Se concluye que el enfoque permite futuras simulaciones de tribunal. "
        )
        * 35
    )
    document.save(buffer)
    return buffer.getvalue()


def register(client: TestClient, email: str) -> None:
    response = client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "full_name": "Question Bank Student",
            "password": "secure-question-bank-password",
        },
    )
    assert response.status_code == 201


def upload_document(client: TestClient) -> str:
    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "defensa.docx",
                valid_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def processed_document(client: TestClient, email: str) -> str:
    register(client, email)
    document_id = upload_document(client)
    response = client.post(f"/api/v1/documents/{document_id}/process")
    assert response.status_code == 200
    return document_id


def provider_failure(
    provider: str,
    kind: LLMErrorKind,
    detail: str | None = None,
) -> QuestionGenerationProviderError:
    return QuestionGenerationProviderError(
        "simulated provider failure",
        provider=provider,
        kind=kind,
        detail=detail,
    )


def test_question_bank_requires_authentication(client: TestClient) -> None:
    generate = client.post("/api/v1/documents/not-found/questions/generate")
    get = client.get("/api/v1/documents/not-found/questions")

    assert generate.status_code == 401
    assert get.status_code == 401


def test_question_bank_requires_processed_document(client: TestClient) -> None:
    register(client, "not-processed@example.com")
    document_id = upload_document(client)

    response = client.post(f"/api/v1/documents/{document_id}/questions/generate")

    assert response.status_code == 409
    assert response.json()["error"]["code"] == "document_not_processed"


def test_question_bank_enforces_ownership(client: TestClient) -> None:
    document_id = processed_document(client, "question-owner@example.com")
    register(client, "question-intruder@example.com")

    generate = client.post(f"/api/v1/documents/{document_id}/questions/generate")
    get = client.get(f"/api/v1/documents/{document_id}/questions")

    assert generate.status_code == 404
    assert get.status_code == 404


def test_generation_filters_deduplicates_persists_and_hides_internal_points(
    client: TestClient,
    db_session: Session,
    embedding_provider: Any,
    vector_store: Any,
    question_primary: Any,
) -> None:
    document_id = processed_document(client, "rag-bank@example.com")

    response = client.post(f"/api/v1/documents/{document_id}/questions/generate")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "READY"
    assert payload["provider_used"] == "gemini"
    assert payload["model_used"] == "fake-gemini"
    assert payload["fallback_used"] is False
    assert len(payload["questions"]) == 12
    assert all("expected_answer_points" not in item for item in payload["questions"])
    assert all("source_chunk_ids" not in item for item in payload["questions"])
    assert {category: sum(q["category"] == category for q in payload["questions"]) for category in (
        "CONCEPTUAL",
        "METHODOLOGICAL",
        "TECHNICAL",
        "CRITICAL",
    )} == {
        "CONCEPTUAL": 3,
        "METHODOLOGICAL": 3,
        "TECHNICAL": 3,
        "CRITICAL": 3,
    }

    assert len(embedding_provider.queries) == 6
    assert len(vector_store.query_calls) == 6
    for query_call in vector_store.query_calls:
        assert query_call["top_k"] == 4
        assert query_call["filters"] == {
            "$and": [
                {"user_id": {"$eq": payload_document_user(db_session, document_id)}},
                {"document_id": {"$eq": document_id}},
            ]
        }
    request = question_primary.requests[0]
    context_ids = [chunk.id for chunk in request.chunks]
    assert len(context_ids) == len(set(context_ids))

    assert db_session.scalar(select(func.count()).select_from(QuestionBank)) == 1
    assert db_session.scalar(select(func.count()).select_from(Question)) == 12
    stored_question = db_session.scalars(select(Question).order_by(Question.position)).first()
    assert stored_question is not None
    assert len(stored_question.expected_answer_points) == 2

    cached = client.get(f"/api/v1/documents/{document_id}/questions")
    assert cached.status_code == 200
    assert cached.json()["id"] == payload["id"]
    assert len(question_primary.requests) == 1


def payload_document_user(db_session: Session, document_id: str) -> str:
    bank = db_session.scalar(select(QuestionBank).where(QuestionBank.document_id == document_id))
    assert bank is not None
    return bank.user_id


@pytest.mark.parametrize(
    ("failure_kind", "detail", "invalid_source"),
    [
        (LLMErrorKind.TRANSIENT, "network", False),
        (LLMErrorKind.TRANSIENT, "http_500", False),
        (LLMErrorKind.INVALID_OUTPUT, None, False),
        (None, None, True),
    ],
)
def test_generation_falls_back_to_groq(
    client: TestClient,
    question_primary: Any,
    question_fallback: Any,
    failure_kind: LLMErrorKind | None,
    detail: str | None,
    invalid_source: bool,
) -> None:
    document_id = processed_document(client, f"fallback-{detail or 'invalid-source'}@example.com")
    if failure_kind is not None:
        question_primary.failure = provider_failure("gemini", failure_kind, detail)
    question_primary.invalid_source = invalid_source

    response = client.post(f"/api/v1/documents/{document_id}/questions/generate")

    assert response.status_code == 200
    payload = response.json()
    assert payload["provider_used"] == "groq"
    assert payload["model_used"] == "fake-groq"
    assert payload["fallback_used"] is True
    assert len(question_fallback.requests) == 1


def test_all_provider_failures_preserve_processed_document(
    client: TestClient,
    db_session: Session,
    question_primary: Any,
    question_fallback: Any,
) -> None:
    document_id = processed_document(client, "both-providers-fail@example.com")
    question_primary.failure = provider_failure("gemini", LLMErrorKind.TRANSIENT, "network")
    question_fallback.failure = provider_failure("groq", LLMErrorKind.TRANSIENT, "http_503")

    response = client.post(f"/api/v1/documents/{document_id}/questions/generate")

    assert response.status_code == 503
    assert response.json()["error"]["code"] == "QUESTION_GENERATION_FAILED"
    assert client.get(f"/api/v1/documents/{document_id}/status").json()["status"] == "PROCESSED"
    bank = db_session.scalar(select(QuestionBank).where(QuestionBank.document_id == document_id))
    assert bank is not None
    assert bank.status == QuestionBankStatus.FAILED
    assert "gemini:transient:network" in (bank.failure_reason or "")
    assert "groq:transient:http_503" in (bank.failure_reason or "")


def test_get_does_not_generate_and_regeneration_replaces_questions(
    client: TestClient,
    db_session: Session,
    question_primary: Any,
) -> None:
    document_id = processed_document(client, "regenerate@example.com")
    missing = client.get(f"/api/v1/documents/{document_id}/questions")
    assert missing.status_code == 404
    assert len(question_primary.requests) == 0

    first = client.post(f"/api/v1/documents/{document_id}/questions/generate").json()
    first_question_ids = {question["id"] for question in first["questions"]}
    second = client.post(f"/api/v1/documents/{document_id}/questions/generate").json()

    assert second["id"] == first["id"]
    assert first_question_ids.isdisjoint({question["id"] for question in second["questions"]})
    assert db_session.scalar(select(func.count()).select_from(QuestionBank)) == 1
    assert db_session.scalar(select(func.count()).select_from(Question)) == 12
    assert len(question_primary.requests) == 2


def test_delete_document_cascades_question_bank(
    client: TestClient,
    db_session: Session,
) -> None:
    document_id = processed_document(client, "delete-bank@example.com")
    assert client.post(f"/api/v1/documents/{document_id}/questions/generate").status_code == 200

    response = client.delete(f"/api/v1/documents/{document_id}")

    assert response.status_code == 204
    assert db_session.scalar(select(func.count()).select_from(QuestionBank)) == 0
    assert db_session.scalar(select(func.count()).select_from(Question)) == 0


def test_circuit_breaker_skips_open_primary(
    question_primary: Any,
    question_fallback: Any,
) -> None:
    question_primary.failure = provider_failure("gemini", LLMErrorKind.TRANSIENT, "network")
    router = QuestionGenerationRouter(
        primary=question_primary,
        fallback=question_fallback,
        failure_threshold=1,
        recovery_seconds=60,
    )
    request = QuestionGenerationRequest(
        document_name="test.docx",
        chunks=(QuestionContextChunk(id="doc:0", text="Contexto academico valido.", score=1.0),),
    )

    first = router.generate(request)
    second = router.generate(request)

    assert first.fallback_used is True
    assert second.fallback_used is True
    assert second.primary_failure_reason == "gemini:circuit_open"
    assert len(question_primary.requests) == 1
    assert len(question_fallback.requests) == 2
