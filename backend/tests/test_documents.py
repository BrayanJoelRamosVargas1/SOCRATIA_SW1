from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.integrations.embeddings.base import EmbeddingDocument
from app.integrations.storage.local import LocalStorageProvider
from app.integrations.vector_db.base import VectorRecord
from app.modules.p2_gestion_documentos_preparacion.models.document import (
    Document as StoredDocument,
)
from app.modules.p2_gestion_documentos_preparacion.models.document import DocumentStatus

PDF_CONTENT = b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\n%%EOF"


def valid_docx(text: str = "Socratia document content") -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_heading("Socratia", level=1)
    document.add_paragraph(text)
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "Academic material"
    document.save(buffer)
    return buffer.getvalue()


def empty_docx() -> bytes:
    buffer = BytesIO()
    DocxDocument().save(buffer)
    return buffer.getvalue()


def valid_pdf(text: str = "Socratia academic document") -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{number} 0 obj\n".encode())
        result.extend(body)
        result.extend(b"\nendobj\n")
    xref_offset = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode())
    result.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF"
        ).encode()
    )
    return bytes(result)


def register(client: TestClient, email: str) -> None:
    response = client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "full_name": "Document Student",
            "password": "secure-document-password",
        },
    )
    assert response.status_code == 201


def upload_pdf(client: TestClient, filename: str = "tesis.pdf") -> dict:
    response = client.post(
        "/api/v1/documents",
        files={"file": (filename, PDF_CONTENT, "application/pdf")},
    )
    assert response.status_code == 201
    return response.json()


def stored_files(storage: LocalStorageProvider) -> list[Path]:
    if not storage.root.exists():
        return []
    return [path for path in storage.root.rglob("*") if path.is_file()]


def test_upload_document(
    client: TestClient,
    storage_provider: LocalStorageProvider,
) -> None:
    register(client, "upload@example.com")

    document = upload_pdf(client)

    assert document["original_name"] == "tesis.pdf"
    assert document["file_type"] == "PDF"
    assert document["file_size"] == len(PDF_CONTENT)
    assert document["status"] == "UPLOADED"
    files = stored_files(storage_provider)
    assert len(files) == 1
    assert files[0].read_bytes() == PDF_CONTENT


def test_list_own_documents(client: TestClient) -> None:
    register(client, "list@example.com")
    first = upload_pdf(client, "tesis.pdf")
    second = upload_pdf(client, "articulo.pdf")

    response = client.get("/api/v1/documents")

    assert response.status_code == 200
    assert {item["id"] for item in response.json()} == {first["id"], second["id"]}


def test_upload_docx(client: TestClient) -> None:
    register(client, "docx@example.com")
    content = valid_docx()

    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "research.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    assert response.json()["file_type"] == "DOCX"
    assert response.json()["file_size"] == len(content)


def test_cannot_read_other_user_document(client: TestClient) -> None:
    register(client, "owner@example.com")
    document = upload_pdf(client)
    register(client, "intruder@example.com")

    detail = client.get(f"/api/v1/documents/{document['id']}")
    processing = client.get(f"/api/v1/documents/{document['id']}/status")
    deletion = client.delete(f"/api/v1/documents/{document['id']}")

    assert detail.status_code == 404
    assert processing.status_code == 404
    assert deletion.status_code == 404
    assert client.get("/api/v1/documents").json() == []


def test_get_processing_status(client: TestClient) -> None:
    register(client, "status@example.com")
    document = upload_pdf(client)

    response = client.get(f"/api/v1/documents/{document['id']}/status")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "UPLOADED"
    assert payload["chunk_count"] == 0
    assert len(payload["history"]) == 1
    assert payload["history"][0]["stage"] == "UPLOAD"
    assert payload["history"][0]["status"] == "UPLOADED"


def test_reject_invalid_file(client: TestClient) -> None:
    register(client, "invalid@example.com")

    wrong_extension = client.post(
        "/api/v1/documents",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    spoofed_pdf = client.post(
        "/api/v1/documents",
        files={"file": ("fake.pdf", b"not a pdf", "application/pdf")},
    )

    assert wrong_extension.status_code == 422
    assert wrong_extension.json()["error"]["code"] == "invalid_document_file"
    assert spoofed_pdf.status_code == 422
    assert client.get("/api/v1/documents").json() == []


def test_delete_document(
    client: TestClient,
    storage_provider: LocalStorageProvider,
) -> None:
    register(client, "delete@example.com")
    document = upload_pdf(client)
    assert len(stored_files(storage_provider)) == 1

    response = client.delete(f"/api/v1/documents/{document['id']}")

    assert response.status_code == 204
    assert stored_files(storage_provider) == []
    assert client.get(f"/api/v1/documents/{document['id']}").status_code == 404


def test_process_docx_with_embeddings_and_vectors(
    client: TestClient,
    embedding_provider: Any,
    vector_store: Any,
) -> None:
    register(client, "processing-docx@example.com")
    content = valid_docx("Socratic learning strengthens critical thinking. " * 120)
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "research.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document = upload.json()

    response = client.post(f"/api/v1/documents/{document['id']}/process")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "PROCESSED"
    assert payload["chunk_count"] >= 2
    assert payload["embedding_model"] == "fake-embedding"
    assert payload["embedding_dimensions"] == 3
    assert len(embedding_provider.documents) == payload["chunk_count"]
    assert all(isinstance(item, EmbeddingDocument) for item in embedding_provider.documents)
    records = next(iter(vector_store.namespaces.values()))
    assert len(records) == payload["chunk_count"]
    assert all(isinstance(record, VectorRecord) for record in records.values())
    assert all(record.metadata["document_id"] == document["id"] for record in records.values())

    processing_status = client.get(f"/api/v1/documents/{document['id']}/status").json()
    assert processing_status["status"] == "PROCESSED"
    assert processing_status["chunk_count"] == payload["chunk_count"]
    assert [step["stage"] for step in processing_status["history"]][-5:] == [
        "EXTRACTION",
        "CHUNKING",
        "EMBEDDING",
        "VECTOR_STORE",
        "COMPLETE",
    ]


def test_process_pdf(client: TestClient) -> None:
    register(client, "processing-pdf@example.com")
    content = valid_pdf("Socratia extracts text from a real PDF document.")
    upload = client.post(
        "/api/v1/documents",
        files={"file": ("paper.pdf", content, "application/pdf")},
    )

    response = client.post(f"/api/v1/documents/{upload.json()['id']}/process")

    assert response.status_code == 200
    assert response.json()["status"] == "PROCESSED"
    assert response.json()["chunk_count"] == 1


def test_processing_provider_failure_is_recorded(
    client: TestClient,
    embedding_provider: Any,
) -> None:
    register(client, "processing-error@example.com")
    content = valid_docx("A valid document whose embedding request will fail.")
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "failure.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = upload.json()["id"]
    embedding_provider.fail = True

    response = client.post(f"/api/v1/documents/{document_id}/process")

    assert response.status_code == 503
    assert response.json()["error"]["code"] == "document_processing_unavailable"
    processing_status = client.get(f"/api/v1/documents/{document_id}/status").json()
    assert processing_status["status"] == "ERROR"
    assert processing_status["history"][-1]["stage"] == "EMBEDDING"
    assert processing_status["history"][-1]["status"] == "ERROR"


def test_cannot_process_other_user_document(client: TestClient) -> None:
    register(client, "processing-owner@example.com")
    content = valid_docx("Private academic content")
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "private.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    register(client, "processing-intruder@example.com")

    response = client.post(f"/api/v1/documents/{upload.json()['id']}/process")

    assert response.status_code == 404


def test_delete_processed_document_removes_vectors(
    client: TestClient,
    vector_store: Any,
) -> None:
    register(client, "processed-delete@example.com")
    content = valid_docx("Document that will be processed and deleted.")
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "delete.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = upload.json()["id"]
    assert client.post(f"/api/v1/documents/{document_id}/process").status_code == 200
    assert any(vector_store.namespaces.values())

    response = client.delete(f"/api/v1/documents/{document_id}")

    assert response.status_code == 204
    assert all(not records for records in vector_store.namespaces.values())


def test_cannot_delete_document_while_processing(
    client: TestClient,
    db_session: Session,
    storage_provider: LocalStorageProvider,
) -> None:
    register(client, "processing-delete@example.com")
    document = upload_pdf(client)
    stored_document = db_session.get(StoredDocument, document["id"])
    assert stored_document is not None
    stored_document.status = DocumentStatus.PROCESSING
    db_session.commit()

    response = client.delete(f"/api/v1/documents/{document['id']}")

    assert response.status_code == 409
    assert response.json()["error"]["code"] == "document_already_processing"
    assert len(stored_files(storage_provider)) == 1


def test_reprocess_document_reuses_deterministic_vector_ids(
    client: TestClient,
    vector_store: Any,
) -> None:
    register(client, "processing-retry@example.com")
    content = valid_docx("Repeatable academic content. " * 150)
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "retry.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = upload.json()["id"]

    first = client.post(f"/api/v1/documents/{document_id}/process")
    first_ids = set(next(iter(vector_store.namespaces.values())))
    second = client.post(f"/api/v1/documents/{document_id}/process")
    second_ids = set(next(iter(vector_store.namespaces.values())))

    assert first.status_code == 200
    assert second.status_code == 200
    assert second.json()["chunk_count"] == first.json()["chunk_count"]
    assert second_ids == first_ids


def test_reject_document_without_extractable_text(client: TestClient) -> None:
    register(client, "processing-empty@example.com")
    upload = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "empty.docx",
                empty_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = upload.json()["id"]

    response = client.post(f"/api/v1/documents/{document_id}/process")

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "document_content_unreadable"
    processing_status = client.get(f"/api/v1/documents/{document_id}/status").json()
    assert processing_status["status"] == "ERROR"
    assert processing_status["history"][-1]["stage"] == "EXTRACTION"


def test_reject_second_processing_while_active(
    client: TestClient,
    db_session: Session,
) -> None:
    register(client, "processing-active@example.com")
    document = upload_pdf(client)
    stored_document = db_session.get(StoredDocument, document["id"])
    assert stored_document is not None
    stored_document.status = DocumentStatus.PROCESSING
    db_session.commit()

    response = client.post(f"/api/v1/documents/{document['id']}/process")

    assert response.status_code == 409
    assert response.json()["error"]["code"] == "document_already_processing"
