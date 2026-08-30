from io import BytesIO
from typing import Any

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.integrations.llm import LLMErrorKind, QuestionGenerationProviderError
from app.modules.p2_gestion_documentos_preparacion.models.presentation_material import (
    PresentationMaterial,
    PresentationMaterialStatus,
    PresentationSlide,
)


def valid_docx() -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    for title, content in (
        ("Problema", "Las defensas academicas carecen de preparacion guiada."),
        ("Objetivos", "Preparar al estudiante con evidencia de su propio documento."),
        ("Metodologia", "Se aplica recuperacion semantica aislada por usuario."),
        ("Desarrollo", "La arquitectura integra FastAPI, PostgreSQL y Pinecone."),
        ("Resultados", "Las pruebas validan trazabilidad, seguridad y resiliencia."),
        ("Conclusiones", "El sistema permite practicar una defensa sustentada."),
    ):
        document.add_heading(title, level=1)
        document.add_paragraph((content + " ") * 80)
    document.save(buffer)
    return buffer.getvalue()


def register(client: TestClient, email: str) -> None:
    response = client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "full_name": "Presentation Student",
            "password": "secure-presentation-password",
        },
    )
    assert response.status_code == 201


def upload(client: TestClient) -> str:
    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "defensa.docx",
                valid_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def processed(client: TestClient, email: str) -> str:
    register(client, email)
    document_id = upload(client)
    assert client.post(f"/api/v1/documents/{document_id}/process").status_code == 200
    return document_id


def failure(provider: str) -> QuestionGenerationProviderError:
    return QuestionGenerationProviderError(
        "simulated",
        provider=provider,
        kind=LLMErrorKind.TRANSIENT,
        detail="network",
    )


def test_presentation_requires_authentication(client: TestClient) -> None:
    assert client.post(
        "/api/v1/documents/missing/presentation/generate",
        json={"duration_minutes": 15},
    ).status_code == 401
    assert client.get("/api/v1/documents/missing/presentation").status_code == 401


@pytest.mark.parametrize("duration", [0, 4, 31, 900])
def test_presentation_rejects_duration_outside_mvp_range(
    client: TestClient, duration: int
) -> None:
    register(client, f"duration-{duration}@example.com")
    response = client.post(
        "/api/v1/documents/missing/presentation/generate",
        json={"duration_minutes": duration},
    )
    assert response.status_code == 422


def test_presentation_requires_processed_owned_document(client: TestClient) -> None:
    register(client, "unprocessed-presentation@example.com")
    document_id = upload(client)
    response = client.post(
        f"/api/v1/documents/{document_id}/presentation/generate",
        json={"duration_minutes": 15},
    )
    assert response.status_code == 409

    processed_id = processed(client, "presentation-owner@example.com")
    register(client, "presentation-intruder@example.com")
    assert client.post(
        f"/api/v1/documents/{processed_id}/presentation/generate",
        json={"duration_minutes": 15},
    ).status_code == 404
    assert client.get(f"/api/v1/documents/{processed_id}/presentation").status_code == 404


def test_presentation_retrieves_persists_and_hides_source_chunks(
    client: TestClient,
    db_session: Session,
    embedding_provider: Any,
    vector_store: Any,
    presentation_primary: Any,
) -> None:
    document_id = processed(client, "presentation-rag@example.com")
    before_queries = len(vector_store.query_calls)
    response = client.post(
        f"/api/v1/documents/{document_id}/presentation/generate",
        json={"duration_minutes": 15},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "READY"
    assert payload["provider_used"] == "gemini"
    assert payload["duration_minutes"] == 15
    assert payload["target_word_count"] == 1950
    assert 8 <= len(payload["slides"]) <= 12
    assert sum(slide["estimated_seconds"] for slide in payload["slides"]) == 900
    assert all("source_chunk_ids" not in slide for slide in payload["slides"])
    assert len(embedding_provider.queries) >= 6
    calls = vector_store.query_calls[before_queries:]
    assert len(calls) == 6
    for call in calls:
        clauses = call["filters"]["$and"]
        assert {"document_id": {"$eq": document_id}} in clauses
        assert any("user_id" in clause for clause in clauses)

    assert db_session.scalar(select(func.count()).select_from(PresentationMaterial)) == 1
    assert db_session.scalar(select(func.count()).select_from(PresentationSlide)) == len(
        payload["slides"]
    )
    stored = db_session.scalars(select(PresentationSlide)).first()
    assert stored is not None and stored.source_chunk_ids

    cached = client.get(f"/api/v1/documents/{document_id}/presentation")
    assert cached.status_code == 200
    assert cached.json()["id"] == payload["id"]
    assert len(presentation_primary.requests) == 1


def test_generate_conflicts_and_regenerate_replaces_material(
    client: TestClient,
    db_session: Session,
) -> None:
    document_id = processed(client, "presentation-regenerate@example.com")
    url = f"/api/v1/documents/{document_id}/presentation"
    first = client.post(f"{url}/generate", json={"duration_minutes": 15}).json()
    first_slide_ids = {slide["id"] for slide in first["slides"]}

    duplicate = client.post(f"{url}/generate", json={"duration_minutes": 10})
    assert duplicate.status_code == 409
    second = client.post(f"{url}/regenerate", json={"duration_minutes": 10})
    assert second.status_code == 200
    payload = second.json()
    assert payload["id"] == first["id"]
    assert payload["duration_minutes"] == 10
    assert first_slide_ids.isdisjoint({slide["id"] for slide in payload["slides"]})
    assert db_session.scalar(select(func.count()).select_from(PresentationMaterial)) == 1


def test_invalid_primary_output_falls_back_to_groq(
    client: TestClient,
    presentation_primary: Any,
    presentation_fallback: Any,
) -> None:
    document_id = processed(client, "presentation-fallback@example.com")
    presentation_primary.invalid_source = True
    response = client.post(
        f"/api/v1/documents/{document_id}/presentation/generate",
        json={"duration_minutes": 15},
    )
    assert response.status_code == 200
    assert response.json()["provider_used"] == "groq"
    assert response.json()["fallback_used"] is True
    assert len(presentation_fallback.requests) == 1


def test_both_provider_failures_preserve_document_and_failed_state(
    client: TestClient,
    db_session: Session,
    presentation_primary: Any,
    presentation_fallback: Any,
) -> None:
    document_id = processed(client, "presentation-failed@example.com")
    presentation_primary.failure = failure("gemini")
    presentation_fallback.failure = failure("groq")
    response = client.post(
        f"/api/v1/documents/{document_id}/presentation/generate",
        json={"duration_minutes": 15},
    )
    assert response.status_code == 503
    assert response.json()["error"]["code"] == "PRESENTATION_GENERATION_FAILED"
    assert client.get(f"/api/v1/documents/{document_id}/status").json()["status"] == "PROCESSED"
    material = db_session.scalar(
        select(PresentationMaterial).where(PresentationMaterial.document_id == document_id)
    )
    assert material is not None and material.status == PresentationMaterialStatus.FAILED


def test_delete_document_cascades_presentation(client: TestClient, db_session: Session) -> None:
    document_id = processed(client, "presentation-delete@example.com")
    assert client.post(
        f"/api/v1/documents/{document_id}/presentation/generate",
        json={"duration_minutes": 5},
    ).status_code == 200
    assert client.delete(f"/api/v1/documents/{document_id}").status_code == 204
    assert db_session.scalar(select(func.count()).select_from(PresentationMaterial)) == 0
    assert db_session.scalar(select(func.count()).select_from(PresentationSlide)) == 0
