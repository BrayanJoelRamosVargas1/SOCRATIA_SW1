from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.modules.p3_gestion_simulacion.exceptions import InvalidSimulationTransitionError
from app.modules.p3_gestion_simulacion.models import (
    JuryProfile,
    Simulation,
    SimulationQuestion,
    SimulationStatus,
)
from app.modules.p3_gestion_simulacion.services.simulation_service import SimulationService


def register(client: TestClient, email: str) -> None:
    response = client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "full_name": "Simulation Student",
            "password": "secure-simulation-password",
        },
    )
    assert response.status_code == 201


def document_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Defensa Socratia", level=1)
    document.add_paragraph(
        (
            "Problema objetivos metodologia arquitectura resultados limitaciones conclusiones. "
            "La evidencia sustenta preguntas de defensa y decisiones tecnicas. "
        )
        * 120
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def upload(client: TestClient) -> str:
    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "simulacion.docx",
                document_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def ready_document(client: TestClient, email: str) -> str:
    register(client, email)
    document_id = upload(client)
    assert client.post(f"/api/v1/documents/{document_id}/process").status_code == 200
    assert client.post(f"/api/v1/documents/{document_id}/questions/generate").status_code == 200
    return document_id


def profiles(client: TestClient) -> list[dict]:
    response = client.get("/api/v1/jury-profiles")
    assert response.status_code == 200
    return response.json()


def create(client: TestClient, document_id: str, jury_id: str, duration: int = 15):
    return client.post(
        "/api/v1/simulations",
        json={
            "document_id": document_id,
            "jury_profile_id": jury_id,
            "planned_duration_minutes": duration,
        },
    )


def test_simulation_routes_require_authentication(client: TestClient) -> None:
    assert client.get("/api/v1/jury-profiles").status_code == 401
    assert client.get("/api/v1/simulations").status_code == 401
    assert client.post("/api/v1/simulations", json={}).status_code == 401


def test_default_jury_profiles_are_seeded_idempotently(
    client: TestClient, db_session: Session
) -> None:
    register(client, "jury-profiles@example.com")
    first = profiles(client)
    second = profiles(client)
    assert [item["name"] for item in first] == [
        "Jurado Metodológico",
        "Jurado Técnico",
        "Jurado Crítico",
    ]
    assert first == second
    assert db_session.scalar(select(func.count()).select_from(JuryProfile)) == 3


def test_create_requires_processed_document_and_ready_bank(client: TestClient) -> None:
    register(client, "simulation-prerequisites@example.com")
    jury_id = profiles(client)[0]["id"]
    document_id = upload(client)
    assert create(client, document_id, jury_id).status_code == 409
    assert client.post(f"/api/v1/documents/{document_id}/process").status_code == 200
    assert create(client, document_id, jury_id).status_code == 409


def test_create_rejects_foreign_document(client: TestClient) -> None:
    document_id = ready_document(client, "simulation-owner@example.com")
    register(client, "simulation-intruder@example.com")
    jury_id = profiles(client)[0]["id"]
    assert create(client, document_id, jury_id).status_code == 404


def test_create_freezes_questions_and_lists_only_owner(
    client: TestClient, db_session: Session
) -> None:
    document_id = ready_document(client, "simulation-create@example.com")
    jury = profiles(client)[1]
    response = create(client, document_id, jury["id"])
    assert response.status_code == 201
    payload = response.json()
    assert payload["status"] == "DRAFT"
    assert payload["question_count"] == 12
    assert payload["document"]["id"] == document_id
    assert payload["jury_profile"]["name"] == "Jurado Técnico"
    assert db_session.scalar(select(func.count()).select_from(Simulation)) == 1
    assert db_session.scalar(select(func.count()).select_from(SimulationQuestion)) == 12
    assert client.get("/api/v1/simulations").json()[0]["id"] == payload["id"]
    assert client.get(f"/api/v1/simulations/{payload['id']}").status_code == 200

    register(client, "simulation-other-list@example.com")
    assert client.get("/api/v1/simulations").json() == []
    assert client.get(f"/api/v1/simulations/{payload['id']}").status_code == 404


def test_calibration_persists_only_readiness_and_moves_to_ready(client: TestClient) -> None:
    document_id = ready_document(client, "simulation-calibration@example.com")
    simulation = create(client, document_id, profiles(client)[2]["id"]).json()
    url = f"/api/v1/simulations/{simulation['id']}/calibration"
    partial = client.put(
        url,
        json={"camera_ready": True, "microphone_ready": True, "vision_ready": False},
    )
    assert partial.status_code == 200
    assert partial.json()["status"] == "DRAFT"
    ready = client.put(
        url,
        json={"camera_ready": True, "microphone_ready": True, "vision_ready": True},
    )
    assert ready.status_code == 200
    assert ready.json()["status"] == "READY"
    assert ready.json()["camera_ready"] is True
    assert "frame" not in ready.json() and "audio" not in ready.json()
    assert client.put(
        url,
        json={
            "camera_ready": True,
            "microphone_ready": True,
            "vision_ready": True,
            "video": "never accepted",
        },
    ).status_code == 422


def test_delete_only_draft_or_ready_and_state_machine_rejects_terminal_transition(
    client: TestClient, db_session: Session
) -> None:
    document_id = ready_document(client, "simulation-delete@example.com")
    simulation_id = create(client, document_id, profiles(client)[0]["id"]).json()["id"]
    assert client.delete(f"/api/v1/simulations/{simulation_id}").status_code == 204

    terminal = Simulation(
        user_id="u",
        document_id="d",
        question_bank_id="b",
        jury_profile_id="j",
        status=SimulationStatus.COMPLETED,
        planned_duration_minutes=15,
    )
    try:
        SimulationService(db_session).transition(terminal, SimulationStatus.ACTIVE)
    except InvalidSimulationTransitionError:
        pass
    else:
        raise AssertionError("COMPLETED -> ACTIVE must be rejected")
